import os
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
import datetime
import argparse

from analysis_asm import (
    run_audit,
    compliance,
    evasion,
    cookies,
    filetypes,
    headers,
    parameters,
    urls,
    signatures_summary,
    list_to_dict,
    dict_to_list,
)  # your ASM run_audit and helpers


# -------------------------------------------------------------------
# Helper functions
# -------------------------------------------------------------------

def safe_get(value, default="not configured"):
    if value is None or value == "" or str(value).strip() == "":
        return default
    return str(value)


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def sanitize_filename(name: str) -> str:
    """
    Make a filesystem-safe version of the policy name.
    """
    if not name:
        return "Unnamed_Policy"
    return "".join(
        c if c.isalnum() or c in (" ", "-", "_") else "_"
        for c in name
    ).strip() or "Unnamed_Policy"


# -------------------------------------------------------------------
# Report generation for a single policy
# -------------------------------------------------------------------

def generate_report_for_policy(policy: dict, customer_name: str, report_date: datetime.date) -> Path:
    """
    Generate a Word report for a single ASM policy.
    Returns the Path of the saved report.
    """
    policy_name = policy.get("name", "Unnamed Policy")
    today_str = report_date.isoformat()

    document = Document("Template.docx")

    # ---------- Title & Intro ----------

    document.add_heading("F5 ASM Configuration Review - " + policy_name, level=1)
    # Explicit date line
    document.add_paragraph(f"Report date: {today_str}")

    document.add_paragraph(
        'During the meetings with "' + customer_name +
        '" and the configuration reviews performed on the F5 BIG-IP ASM policies, '
        "we documented our observations and recommendations in this report."
    )

    document.add_paragraph(
        "This document presents the findings of the F5 BIG-IP Application Security Manager (ASM) "
        f'configuration review performed for the application "{policy_name}". '
        "The purpose of this review is to assess the current web application firewall policies, "
        "identify potential gaps, and verify alignment with F5 recommended practices and common "
        "security standards."
    )

    document.add_paragraph(
        "The focus of this review is on the security policy configuration, including signature usage, "
        "protocol and evasion checks, cookie and parameter protection, URL handling, brute-force "
        "protection, IP reputation, and related security controls. The goal is to highlight areas "
        "where the current configuration may allow unnecessary risk, be overly permissive, or require "
        "tuning to better balance security with application behaviour."
    )

    document.add_paragraph(
        "This report does not evaluate the internal functionality of the protected applications, but "
        "rather how ASM is configured to protect them. The recommendations are intended to help improve "
        "the security posture, reduce the likelihood of successful attacks, and simplify ongoing operations."
    )

    document.add_paragraph()
    document.add_paragraph()

    # -------------------------------------------------------------------
    # Policy Overview section
    # -------------------------------------------------------------------

    document.add_heading("Policy Overview", level=2)
    document.add_paragraph(
        "The table below provides a high-level summary of the ASM policy configuration that was reviewed."
    )

    applicationLanguage = policy.get("applicationLanguage", "N/A")
    case_insensitive = "Yes" if policy.get("caseInsensitive", True) else "No"
    maskCreditCardNumbersInRequest = "Enabled" if policy.get("general", {}).get("maskCreditCardNumbersInRequest", False) else "Disabled"
    enforcementReadinessPeriod = policy.get("general", {}).get("enforcementReadinessPeriod", 0)
    triggerAsmIruleEvent = policy.get("general", {}).get("triggerAsmIruleEvent", "N/A")
    trust_all_ips = "Yes" if policy.get("policy-builder", {}).get("trustAllIps", False) else "No"
    trustXff = "Enabled" if policy.get("general", {}).get("trustXff", False) else "Disabled"
    min_accuracy = policy.get("signature-settings", {}).get("minimumAccuracyForAutoAddedSignatures", "Not set")
    place_in_staging = "Enabled" if policy.get("signature-settings", {}).get("placeSignaturesInStaging", False) else "Disabled"

    result_evasion = evasion(policy)
    result_compliance = compliance(policy)
    result_cookies = cookies(policy)
    result_filetypes = filetypes(policy)
    result_headers = headers(policy)
    result_parameters = parameters(policy)
    result_urls = urls(policy)
    result_signatures = signatures_summary(policy)

    builder_settings = policy.get("policy-builder", {})

    table = document.add_table(rows=10, cols=2)
    table.style = "Table Grid"

    cell_merged = table.cell(0, 0).merge(table.cell(0, 1))
    cell_merged.text = "Policy Details"

    table.cell(1, 0).text = "Policy Name"
    table.cell(1, 1).text = policy_name

    table.cell(2, 0).text = "Enforcement Mode"
    table.cell(2, 1).text = policy.get("enforcementMode", "N/A").capitalize()

    table.cell(3, 0).text = "Application Language"
    table.cell(3, 1).text = applicationLanguage

    table.cell(4, 0).text = "Mask credit card numbers in requests"
    table.cell(4, 1).text = maskCreditCardNumbersInRequest

    table.cell(5, 0).text = "Place New/Updated Signatures in Staging"
    table.cell(5, 1).text = place_in_staging

    table.cell(6, 0).text = "Minimum Accuracy for Auto-Added Signatures"
    table.cell(6, 1).text = str(min_accuracy)

    table.cell(7, 0).text = "Trust X-Forwarded-For Header"
    table.cell(7, 1).text = trustXff

    table.cell(8, 0).text = "Enforcement Period"
    table.cell(8, 1).text = str(enforcementReadinessPeriod)

    table.cell(9, 0).text = "Trust all IPs in Policy Builder"
    table.cell(9, 1).text = trust_all_ips

    table = document.add_table(rows=9, cols=3)
    table.style = "Table Grid"

    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("Entities")
    run.bold = True
    table.cell(1, 0).text = "File Types"
    table.cell(2, 0).text = "URLs"
    table.cell(3, 0).text = "Parameters"
    table.cell(4, 0).text = "Signatures"
    table.cell(5, 0).text = "Cookies"
    table.cell(6, 0).text = "Headers"
    table.cell(7, 0).text = "HTTP Compliance"
    table.cell(8, 0).text = "Evasion"

    # ------------------------------------------------------------------
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    run = p.add_run("Total Configured")
    run.bold = True
    table.cell(1, 1).text = str(result_filetypes["total"])
    table.cell(2, 1).text = str(result_urls["total"])
    table.cell(3, 1).text = str(result_parameters["total"])
    table.cell(4, 1).text = str(result_signatures["total"])
    table.cell(5, 1).text = str(result_cookies["total"])
    table.cell(6, 1).text = str(result_headers["total"])
    table.cell(7, 1).text = str(result_compliance["total"])
    table.cell(8, 1).text = str(result_evasion["total"])

    # ------------------------------------------------------------------
    cell = table.cell(0, 2)
    p = cell.paragraphs[0]
    run = p.add_run("Not Enforced (Staging)")
    run.bold = True
    table.cell(1, 2).text = str(result_filetypes["staged"])
    table.cell(2, 2).text = str(result_urls["staged"])
    table.cell(3, 2).text = str(result_parameters["staged"])
    table.cell(4, 2).text = "Staging: " + str(result_signatures["staged"]) + " / " + "Disabled: " + str(result_signatures["disabled"])
    table.cell(5, 2).text = str(result_cookies["staged"])
    table.cell(6, 2).text = str(result_headers["sig_disabled"])
    table.cell(7, 2).text = str(result_compliance["disabled"])
    table.cell(8, 2).text = str(result_evasion["disabled"])

    document.add_paragraph()
    document.add_paragraph()

    # -------------------------------------------------------------------
    # Run ASM Audit and Suggestions / Findings
    # -------------------------------------------------------------------

    suggestions = run_audit(policy)

    document.add_heading("Suggestions and Findings", level=2)
    document.add_paragraph(
        "The following section provides the findings overview from the ASM configuration review."
        "Detail analysis for each finding will follow the subsequent subsections."
    )

    document.add_paragraph()
    document.add_paragraph(
        "The following table provides the summary of all the recommendations of the health check for this ASM policy."
    )

    num_of_suggestions = 0
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.cell(0, 0).text = "#"
    table.cell(0, 1).text = "Severity"
    table.cell(0, 2).text = "Suggestion"
    table.cell(0, 3).text = "Category"

    dict_suggestions = dict_to_list(suggestions)

    # Critical
    for key in dict_suggestions:
        if key["rating"] == "Critical":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.25), height=Inches(.25))
            cells[2].text = key["brief"]
            cells[3].text = key["section"]

    # High
    for key in dict_suggestions:
        if key["rating"] == "High":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.25), height=Inches(.25))
            cells[2].text = key["brief"]
            cells[3].text = key["section"]

    # Medium
    for key in dict_suggestions:
        if key["rating"] == "Medium":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.25), height=Inches(.25))
            cells[2].text = key["title"]
            cells[3].text = key["section"]

    # Low
    for key in dict_suggestions:
        if key["rating"] == "Low":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.25), height=Inches(.25))
            cells[2].text = key["brief"]
            cells[3].text = key["section"]

    # Info
    for key in dict_suggestions:
        if key["rating"] == "Info":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.25), height=Inches(.25))
            cells[2].text = key["brief"]
            cells[3].text = key["section"]

    set_column_width(table.columns[0], Cm(1))
    set_column_width(table.columns[1], Cm(1))
    set_column_width(table.columns[2], Cm(12))
    set_column_width(table.columns[3], Cm(3))

    document.add_paragraph()
    document.add_paragraph()

    document.add_paragraph(
        "The following subsections highlight the main findings from the ASM configuration review."
        "Each item includes a short description, the associated severity, and practical guidance on what should be reviewed or adjusted."
    )

    # --- All the individual suggestion sections (unchanged from your script) ---
    # I’ll keep your existing structure; just re-using `suggestions` as-is

# -------------------------------------------------------------------
# Print findings grouped by severity
# -------------------------------------------------------------------

    SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Info"]

    def add_issue_block(key_name: str):
        issue = suggestions.get(key_name, {})
        if not issue.get("issue"):
            return

        document.add_paragraph()
        document.add_heading(issue["title"], level=3)
        document.add_paragraph(
            "Severity: " + issue["rating"],
            style="sev" + issue["rating"]
        )
        document.add_paragraph(issue["details"])
        document.add_paragraph()
        if issue.get("additional_info"):
            document.add_paragraph(issue["additional_text"])
            additional_info = issue["additional_info"]
            if isinstance(additional_info, list):
                for item in additional_info:
                    document.add_paragraph(f"   -  {item}")
            else:
                document.add_paragraph(str(additional_info))
    # Group keys under each severity
    severity_groups = {sev: [] for sev in SEVERITY_ORDER}

    for key, issue in suggestions.items():
        if issue.get("issue"):
            rating = issue.get("rating", "Info")
            if rating not in severity_groups:
                severity_groups[rating] = []
            severity_groups[rating].append(key)

    # Print in correct order
    for sev in SEVERITY_ORDER:
        for key_name in severity_groups.get(sev, []):
            add_issue_block(key_name)

    # -------------------------------------------------------------------
    # Save report
    # -------------------------------------------------------------------

    output_dir = Path("reports")
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_name = sanitize_filename(policy_name)
    output_file = output_dir / f"F5 ASM - Config Review - {customer_name} - {safe_name} - {today_str}.docx"

    document.save(str(output_file))
    return output_file


# -------------------------------------------------------------------
# Main: read all policies from folder and generate one report per policy
# -------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate ASM reports for all policies in ./policies")
    parser.add_argument(
        "--customer",
        "-c",
        required=True,
        help="Customer name to be used in the report",
    )
    args = parser.parse_args()
    customer_name = args.customer

    policies_dir = Path("policies")
    if not policies_dir.is_dir():
        print(f"❌ Policies folder not found: {policies_dir.resolve()}")
        raise SystemExit(1)

    policy_files = sorted(policies_dir.glob("*.json"))
    if not policy_files:
        print(f"❌ No JSON policy files found in {policies_dir.resolve()}")
        raise SystemExit(1)

    today = datetime.date.today()
    print(f"Found {len(policy_files)} policy file(s) in {policies_dir}.\n")

    for policy_file in policy_files:
        print(f"Processing policy file: {policy_file.name} ...")
        try:
            with policy_file.open("r", encoding="utf-8") as f:
                policy_data = json.load(f)
        except Exception as e:
            print(f"  ⚠️ Error reading {policy_file.name}: {e} (skipping)")
            continue

        # Support both {"policy": {...}} and direct policy at root
        policy = policy_data.get("policy", policy_data)

        if not isinstance(policy, dict):
            print(f"  ⚠️ File {policy_file.name} does not contain a valid policy object (skipping)")
            continue

        output_file = generate_report_for_policy(policy, customer_name, today)
        print(f"  ✅ ASM report saved to: {output_file}")

    print("\nAll done.")


if __name__ == "__main__":
    main()

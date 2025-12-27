
import os
import json
import datetime
import requests
import hashlib
import time
import base64
import re
import argparse
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
from requests.auth import HTTPDigestAuth
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from analysis_ltm import *
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

requests.packages.urllib3.disable_warnings() 


def flatten_ssl_certs_for_excel(ssl_certs_list):
    """
    Flatten ssl_certs so Excel gets 1 row per certificate usage:
      certificate, not_after, days_to_expiry, usage_type, profile, virtual_server
    """
    flat = []
    for cert in ssl_certs_list or []:
        usages = cert.get("usages") or []

        # If there are no usages, still include the certificate as one row
        if not usages:
            flat.append({
                "certificate": cert.get("certificate"),
                "not_after": cert.get("not_after"),
                "days_to_expiry": cert.get("days_to_expiry"),
                "usage_type": "",
                "profile": "",
                "virtual_server": "",
            })
            continue

        for u in usages:
            flat.append({
                "certificate": cert.get("certificate"),
                "not_after": cert.get("not_after"),
                "days_to_expiry": cert.get("days_to_expiry"),
                "usage_type": u.get("type"),
                "profile": u.get("profile"),
                "virtual_server": u.get("virtual_server"),
            })
    return flat

def add_ssl_certs_table_compact_expiring(
    doc: Document,
    ssl_certs: list[dict],
    threshold_days: int = 60
) -> None:
    headers = ["Certificate", "Not After", "Days to Expiry", "Usages"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_column_width(table.columns[0], Cm(6))
    set_column_width(table.columns[1], Cm(2))
    set_column_width(table.columns[2], Cm(1.3))
    set_column_width(table.columns[3], Cm(8))

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    rows_added = 0

    for cert in ssl_certs or []:
        days = cert.get("days_to_expiry")

        # Defensive check (skip missing / non-numeric)
        try:
            days = int(days)
        except (TypeError, ValueError):
            continue

        # 👉 FILTER: only expiring soon
        if days >= threshold_days:
            continue

        row_cells = table.add_row().cells
        rows_added += 1

        row_cells[0].text = cert.get("certificate", "")
        row_cells[1].text = cert.get("not_after", "")
        row_cells[2].text = str(days)

        # Usages (multi-line cell)
        usages_cell = row_cells[3]
        usages_cell.text = ""

        usages = cert.get("usages") or []
        if not usages:
            usages_cell.paragraphs[0].add_run("No usages found")
        else:
            for idx, u in enumerate(usages):
                p = usages_cell.add_paragraph() if idx > 0 else usages_cell.paragraphs[0]
                p.add_run(
                    f"{u.get('type', '')} | "
                    f"{u.get('profile', '')} | "
                    f"VS: {u.get('virtual_server', 'None')}"
                )

    # Optional: if nothing matched, add a single info row
    if rows_added == 0:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[-1])
        row_cells[0].text = f"No certificates expiring in less than {threshold_days} days."

    # Optional: smaller font for long paths
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def load_ignore_list(path: str) -> set:
    ignored = set()
    if not os.path.exists(path):
        return ignored
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            ignored.add(line)
    return ignored

def humanize_int(value: int) -> str:
    """
    Convert an integer to a human-readable format like:
    1200 -> '1.2K', 2000000 -> '2.0M'
    """
    if value is None:
        return "0"

    if value < 1000:
        return str(value)

    for unit in ["K", "M", "G", "T"]:
        value /= 1000.0
        if value < 1000:
            return f"{value:.1f}{unit}"

    # if extremely large (beyond trillions)
    return f"{value:.1f}P"

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

def safe_get(value, default="not configured"):
    if value is None or value == "" or str(value).strip() == "":
        return default
    return str(value)

def normalize_cell(value, none_to_blank=True, list_join="\n"):
    """
    Make values Excel-friendly:
      - None -> ""
      - "none" -> "" (optional)
      - list -> newline-joined
    """    
    if value is None:
        return ""
    if isinstance(value, str):
        if none_to_blank and value.strip().lower() == "none":
            return ""
        return value
    if isinstance(value, list):
        items = [str(x).strip() for x in value if str(x).strip()]
        if none_to_blank and len(items) == 1 and items[0].lower() == "none":
            return ""
        return list_join.join(items)
    return str(value)

def autosize_columns(ws, min_width=12, max_width=80):
    WRAP = Alignment(wrap_text=True, vertical="top")
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            cell.alignment = WRAP
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(min_width, max_len + 2), max_width)

def write_table_sheet(wb, sheet_name, rows, columns, none_to_blank=True):
    """
    Create/replace a sheet and write rows (list[dict]) into it with given columns.
    """

    WRAP = Alignment(wrap_text=True, vertical="top")
    BOLD = Font(bold=True)    
    if sheet_name in wb.sheetnames:
        wb.remove(wb[sheet_name])
    ws = wb.create_sheet(sheet_name)

    # header
    ws.append(columns)
    for i in range(1, len(columns) + 1):
        c = ws.cell(row=1, column=i)
        c.font = BOLD
        c.alignment = WRAP

    # rows
    for r in rows:
        ws.append([normalize_cell(r.get(col), none_to_blank=none_to_blank) for col in columns])

    autosize_columns(ws)

def build_excel(results, customer_name):
    wb = Workbook()   
    # Remove default sheet
    default = wb.active
    wb.remove(default)

    write_table_sheet(
        wb,
        sheet_name="Virtuals",
        rows=results['ltm_virtuals'],
        columns=[
            "virtual",
            "partition",
            "persist",
            "fallback-persistence",
            "policies",
            "profiles",
            "security-log-profiles",
            "rules",
        ],
        none_to_blank=True,
    )

    write_table_sheet(
        wb,
        sheet_name="Pools",
        rows=results['ltm_pools'],
        columns=["pool", "partition", "monitor", "load-balancing-mode"],
        none_to_blank=True,
    )

    write_table_sheet(
        wb,
        sheet_name="HTTP_Profiles",
        rows=results['http'],
        columns=[
            "profile",
            "encrypt-cookies",
            "server-agent-name",
            "insert-xforwarded-for",
            "xff-alternative-names",
        ],
        none_to_blank=True,
    )

    write_table_sheet(
        wb,
        sheet_name="ClientSSL_Profiles",
        rows=results['client_ssl'],
        columns=[
            "profile",
            "cipher-group",
            "ciphers",
            "options",
            "defaults-from",
        ],
        none_to_blank=True,
    )
    write_table_sheet(
        wb,
        sheet_name="iRule Stats",
        rows=results['irule_stats'],
        columns=[
            "name",
            "event",
            "failures",
            "total_executions"
        ],
        none_to_blank=True,
    )
    write_table_sheet(
        wb,
        sheet_name="Virtual_Details",
        rows=results['virtual_details'],
        columns=[
            "virtual",
            "partition",
            "has_dos",
            "has_bot",
            "has_asm",
            "has_client_ssl",
            "has_server_ssl",
            "has_http",
            "has_tcp",
            "has_udp",
            "has_web_acceleration",
            "has_compression",
        ],
        none_to_blank=True,
    )
    write_table_sheet(
        wb,
        sheet_name="VS Traffic",
        rows=results['vs_traffic'],
        columns=["virtual", "clientside_bits_in", "clientside_bits_out", "clientside_packets_in", "clientside_packets_out","clientside_current_connections", "clientside_total_connections", "total_requests", "cpu_usage_ratio_1min"],
        none_to_blank=True,
    )


    ssl_certs_flat = flatten_ssl_certs_for_excel(results['ssl_certs'])
    write_table_sheet(
        wb,
        sheet_name="SSL_Certs",
        rows=ssl_certs_flat,
        columns=[
            "certificate",
            "not_after",
            "days_to_expiry",
            "usage_type",
            "profile",
            "virtual_server",
        ],
        none_to_blank=True,
    )

    today = datetime.date.today()
    today_str = today.isoformat()

    output_dir = Path("reports")
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / f"F5 LTM - Config Review - {customer_name} - {results['list_cm_device']['hostname']} - {today_str}.xlsx"

    wb.save(output_file)
    print(f"Saved Excel: {output_file}")

def build_word(results, diagnostics, customer_name):


    document = Document("Template.docx")

    document.add_heading('F5 LTM Configuration Review', level=1)

    document.add_paragraph('During the meetings with ' + customer_name + ' and the configuration reviews we did on the F5 appliances we present our findings in this report.')

    document.add_paragraph('This document presents the findings of the F5 BIG-IP Local Traffic Manager (LTM) configuration review performed for ' + customer_name + '. The purpose of this review is to provide an in-depth assessment of the current BIG-IP configuration, highlight areas of improvement, and validate whether the deployment follows F5 recommended practices and aligns with industry standards.')

    document.add_paragraph('Throughout the audit, our goal is to help ensure that the BIG-IP platform is operating in a reliable, secure, and efficient manner. By examining the key configuration components of the system, we aim to identify potential misconfigurations, performance bottlenecks, unnecessary complexity, or settings that could introduce operational risk. The outcome of this review is to provide actionable insights that can enhance stability, improve resilience, and optimize traffic management across the environment.')    

    document.add_paragraph('This report also serves as a valuable reference point for future architectural decisions. Whether the customer intends to scale the environment, introduce new services, upgrade software versions, or improve operational processes, the observations in this document provide a baseline understanding of the current state of the configuration. This visibility enables more informed planning and supports long-term maintainability.')    

    document.add_paragraph('While the audit focuses on the technical configuration of the F5 LTM objects, it is not intended to assess application-level logic or behavior.')


    document.add_paragraph()
    document.add_paragraph()

    document.add_heading('Device Details', level=2)
    document.add_paragraph('The following table shows the general configuration from the devices that we have audited. The date and time of the report generation is ' + results["date_info"]["system_date"])

    activeModules = ""
    activeModules2 = text = "\n\n".join(results['license_file']['Active Modules'])
    for key in results['license_file']['Active Modules']:
        activeModules += key
        activeModules += "\n"
        activeModules += "\n"


    table = document.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    input= table.cell(0, 0).merge(table.cell(0, 1))
    input.text = 'Device Details'
    table.cell(1,0).text = 'Platform'
    table.cell(2,0).text = 'Hostname'
    table.cell(3,0).text = 'Version'
    table.cell(4,0).text = 'Chassis ID'
    table.cell(5,0).text = 'DNS Servers'
    table.cell(6,0).text = 'NTP Servers'
    table.cell(7,0).text = 'Time Zone'
    table.cell(8,0).text = 'Licensed Modules'
    #------------------------------------------------------------------
    table.cell(1,1).text = results['sys_hardware']['platform_name']
    table.cell(2,1).text = results['list_cm_device']['hostname']
    table.cell(3,1).text = results['list_cm_device']['version']
    table.cell(4,1).text = results['sys_hardware']['system_chassis_serial']
    table.cell(5,1).text = safe_get(results['sysdb_file'].get('dns.nameservers'))
    table.cell(6,1).text = safe_get(results['sysdb_file'].get('ntp.servers'))
    table.cell(7,1).text = results['list_cm_device']['time-zone']
    table.cell(8,1).text = "\n".join(str(m) for m in results['license_file']['Active Modules']) 


    document.add_heading('System Details', level=2)
    document.add_paragraph('The following table shows the general configuration from the devices that have been configured.')

    size = len(results['sys_provision'])

    table = document.add_table(rows=size+31, cols=2)
    table.style = 'Table Grid'
    input= table.cell(0, 0).merge(table.cell(0, 1))
    input.text = 'Device'
    table.cell(1,0).text = '    Failover State'
    table.cell(2,0).text = '    Active CPUs'
    table.cell(3,0).text = '    Total Memory (MB)'
    table.cell(4,0).text = '    Available Memory (MB)'
    table.cell(1,1).text = str(results['sysdb_file']['failover.state'])
    table.cell(2,1).text = str(results['host_info']['active_cpu_count'])
    table.cell(3,1).text = results['host_info']['memory_total']
    table.cell(4,1).text = results['host_info']['memory_used']
    cell = table.cell(5, 0).merge(table.cell(5, 1))
    # Clear any existing text
    cell.text = ""
    cell.text = ""
    # Add bold text
    p = cell.paragraphs[0]
    run = p.add_run("Configuration Totals")
    run.bold = True
    table.cell(6,0).text = '    Pools'
    table.cell(7,0).text = '    HTTP(s) Monitors'
    table.cell(8,0).text = '    iRules'
    table.cell(9,0).text = '    Certificates'
    table.cell(10,0).text = '    TCP Profiles'
    table.cell(11,0).text = '    UDP Profiles'
    table.cell(12,0).text = '    HTTP Profiles'
    table.cell(13,0).text = '    HTTP2 Profiles'
    table.cell(14,0).text = '    Client SSL Profiles'
    table.cell(15,0).text = '    Server SSL Profiles'
    table.cell(16,0).text = '    Compression Profiles'
    table.cell(17,0).text = '    Web Acceleration Profiles'
    table.cell(18,0).text = '    Bot Policies'
    table.cell(19,0).text = '    Dos Policies'
    table.cell(20,0).text = '    ASM Policies'
    table.cell(6,1).text = str(len(results['ltm_pools']))
    table.cell(7,1).text = str(len(results['http_monitors'])+len(results['https_monitors']))
    table.cell(8,1).text = str(len(results['irule_names']))
    table.cell(9,1).text = str(len(results['ssl_certs']))
    table.cell(10,1).text = str(len(results['tcp_profiles']))
    table.cell(11,1).text = str(len(results['udp_profiles']))
    table.cell(12,1).text = str(len(results['http']))
    table.cell(13,1).text = str(len(results['http2']))
    table.cell(14,1).text = str(len(results['client_ssl']))
    table.cell(15,1).text = str(len(results['server_ssl']))
    table.cell(16,1).text = str(len(results['compression']))
    table.cell(17,1).text = str(len(results['web_acceleration']))
    table.cell(18,1).text = str(len(results['bot']))
    table.cell(19,1).text = str(len(results['dos']))
    table.cell(20,1).text = str(len(results['asm']))
    cell = table.cell(21, 0).merge(table.cell(21, 1))
    # Clear any existing text
    cell.text = ""
    # Add bold text
    p = cell.paragraphs[0]
    run = p.add_run("Virtual Servers Details")
    run.bold = True

    table.cell(22,0).text = '    Virtual Servers Total'
    table.cell(23,0).text = '    L7 Virtual Servers'
    table.cell(24,0).text = '    L4 Virtual Servers '
    table.cell(25,0).text = '    Virtuals with SSL Offloading'
    table.cell(26,0).text = '    Virtuals with Optimization'
    table.cell(27,0).text = '    Virtuals with WAF'
    table.cell(28,0).text = '    Virtuals with Bot'
    table.cell(29,0).text = '    Virtuals with DOS'

    table.cell(22,1).text = str(results['virtual_summary']['count_vs'])
    table.cell(23,1).text = str(results['virtual_summary']['count_vs_http'])
    table.cell(24,1).text = str(results['virtual_summary']['count_vs_tcp']+results['virtual_summary']['count_vs_udp'])
    table.cell(25,1).text = str(results['virtual_summary']['count_vs_client_ssl'])
    table.cell(26,1).text = str(results['virtual_summary']['count_vs_web_acceleration']+results['virtual_summary']['count_vs_compression'])
    table.cell(27,1).text = str(results['virtual_summary']['count_vs_asm'])
    table.cell(28,1).text = str(results['virtual_summary']['count_vs_bot'])
    table.cell(29,1).text = str(results['virtual_summary']['count_vs_dos'])
    cell = table.cell(30, 0).merge(table.cell(30, 1))
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("Provisioned Modules")
    run.bold = True

    start_row = 31
    for i, item in enumerate(results['sys_provision']):
        offset = start_row + i
        table.cell(offset, 0).text = "    " + item["module"].upper()
        table.cell(offset, 1).text = item["level"].capitalize()


    document.add_paragraph()
    document.add_paragraph()
    document.add_heading('Resources Utilization', level=2)
    document.add_paragraph('This section provides a consolidated view of the BIG-IP system’s operational performance across all key hardware and traffic indicators. By visualizing CPU load, memory usage, bandwidth consumption, SSL offloading activity, connection volumes, and HTTP transaction rates, we gain a clear understanding of both platform health and traffic behavior. These insights help identify bottlenecks, verify capacity headroom, and validate that the system is performing reliably under current workloads.')


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("Throughput")
    run.bold = True
    document.add_paragraph('The image below shows the inbound and outbound bandwidth usage.')  
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("HTTP Requests")
    run.bold = True
    document.add_paragraph('The chart below shows the number of HTTP transactions processed.')  
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("Memory Usage")
    run.bold = True
    document.add_paragraph('The chart below displays the memory usage per type.')  
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("SSL Transactions")
    run.bold = True
    document.add_paragraph('The image below shows the device SSL processing levels.')  
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("CPU Utilization")
    run.bold = True
    document.add_paragraph('The image below shows the recorded CPU utilization for the system per CPU core.')
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run("Active Connections")
    run.bold = True
    document.add_paragraph('The image below shows the number of active connections of the device.')  
    p = document.add_paragraph()
    run = p.add_run()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    problems_found = run_audit(results)

    document.add_paragraph()

    document.add_paragraph(
        "The following table provides the summary of all the Issues that have been identified on the BIGIP Device during our audit."
    )

    num_of_suggestions = 0
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "#"
    table.cell(0, 1).text = "Severity"
    table.cell(0, 2).text = "Issue details"

    ignore_set = load_ignore_list("ignore_diagnostics.txt")
    simplified = simplify_diagnostics(diagnostics, ignore_set)
    simplified_non_cve=simplify_non_cve(diagnostics, ignore_set)
    dict_suggestions = dict_to_list(problems_found)


    # Critical
    for key in dict_suggestions:
        if key["rating"] == "Critical":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["title"]



    for key in simplified_non_cve["diagnostics_simplified"]:
        if key["h_importance"] == "CRITICAL":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["h_importance"].capitalize() + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["h_header"]


    # High
    for key in dict_suggestions:
        if key["rating"] == "High":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["title"]
    for key in simplified_non_cve["diagnostics_simplified"]:
        if key["h_importance"] == "HIGH":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["h_importance"].capitalize() + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["h_header"]
    
    if len(simplified["diagnostics_simplified"])>0:
        num_of_suggestions += 1
        cells = table.add_row().cells
        cells[0].text = str(num_of_suggestions)
        paragraph = cells[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture("images/high.png", width=Inches(.18), height=Inches(.18))
        cells[2].text = f"{len(simplified["diagnostics_simplified"])} CVE Vulnerabilities have been identified."

    # Medium
    for key in dict_suggestions:
        if key["rating"] == "Medium":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["title"]
    for key in simplified_non_cve["diagnostics_simplified"]:
        if key["h_importance"] == "MEDIUM":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["h_importance"].capitalize() + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["h_header"]

    # Low
    for key in dict_suggestions:
        if key["rating"] == "Low":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["title"]
    for key in simplified_non_cve["diagnostics_simplified"]:
        if key["h_importance"] == "LOW":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["h_importance"].capitalize() + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["h_header"]

    # Info
    for key in dict_suggestions:
        if key["rating"] == "Info":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["rating"] + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["title"]
    for key in simplified_non_cve["diagnostics_simplified"]:
        if key["h_importance"] == "INFO":
            num_of_suggestions += 1
            cells = table.add_row().cells
            cells[0].text = str(num_of_suggestions)
            paragraph = cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("images/" + key["h_importance"].capitalize() + ".png", width=Inches(.18), height=Inches(.18))
            cells[2].text = key["h_header"]

    set_column_width(table.columns[0], Cm(1))
    set_column_width(table.columns[1], Cm(1))
    set_column_width(table.columns[2], Cm(13))

    document.add_paragraph()

    document.add_paragraph()
    document.add_heading('Health Check Findings', level=2)

    if problems_found['self_ips']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['self_ips']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['self_ips']['rating'], style='sev'+problems_found['self_ips']['rating'])
        document.add_paragraph(problems_found['self_ips']['details'])
        document.add_paragraph()


        document.add_paragraph('Review the configured selfIPs with the Allowed Services per interface.')
        document.add_paragraph()
        size = len(results['self_ips'])
        table = document.add_table(rows=size+1, cols=3)
        table.style = 'Table Grid'
        table.cell(0,0).text = 'Self IP'
        table.cell(0,1).text = 'Address'
        table.cell(0,2).text = 'Allowed Services'
        start_row = 1
        for i, item in enumerate(results['self_ips']):
            offset = start_row + i
            table.cell(offset, 0).text = item['self_ip']
            table.cell(offset, 1).text = item['address']
            table.cell(offset, 2).text = ", ".join(item['allow-service']) if item['allow-service'] else "None"

    if problems_found['ssh_allow']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['ssh_allow']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['ssh_allow']['rating'], style='sev'+problems_found['ssh_allow']['rating'])
        document.add_paragraph(problems_found['ssh_allow']['details'])
        document.add_paragraph()
        document.add_paragraph(problems_found['ssh_allow']['additional'])
        document.add_paragraph()

    if problems_found['httpd_allow']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['httpd_allow']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['httpd_allow']['rating'], style='sev'+problems_found['httpd_allow']['rating'])
        document.add_paragraph(problems_found['httpd_allow']['details'])
        document.add_paragraph()
        document.add_paragraph(problems_found['httpd_allow']['additional'])
        document.add_paragraph()

    if problems_found['password_max']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['password_max']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['password_max']['rating'], style='sev'+problems_found['password_max']['rating'])
        document.add_paragraph(problems_found['password_max']['details'])
        document.add_paragraph()
        document.add_paragraph(problems_found['password_max']['additional'])
        document.add_paragraph()    

    if problems_found['password_remember']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['password_remember']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['password_remember']['rating'], style='sev'+problems_found['password_remember']['rating'])
        document.add_paragraph(problems_found['password_remember']['details'])
        document.add_paragraph()
        document.add_paragraph(problems_found['password_remember']['additional'])
        document.add_paragraph()

    if problems_found['password_complex']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['password_complex']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['password_complex']['rating'], style='sev'+problems_found['password_complex']['rating'])
        document.add_paragraph(problems_found['password_complex']['details'])
        document.add_paragraph()
        document.add_paragraph(problems_found['password_complex']['additional'])
        document.add_paragraph()  

    if problems_found['memory']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['memory']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['memory']['rating'], style='sev'+problems_found['memory']['rating'])
        document.add_paragraph(problems_found['memory']['details'])
        document.add_paragraph()

    if problems_found['sync']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['sync']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['sync']['rating'], style='sev'+problems_found['sync']['rating'])
        document.add_paragraph(problems_found['sync']['details'])
        document.add_paragraph()
        document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
        document.add_paragraph(problems_found['sync']['error'])
        document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    
        document.add_paragraph()

    if problems_found['errors']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['errors']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['errors']['rating'], style='sev'+problems_found['errors']['rating'])
        document.add_paragraph(problems_found['errors']['details'])
        document.add_paragraph()

        document.add_paragraph('Review the interface statistics for the device.')
        document.add_paragraph()
        size = len(results['interface_errors']['interfaces'])
        table = document.add_table(rows=size+1, cols=5)
        table.style = 'Table Grid'

        table.cell(0,0).text = 'Interface'
        table.cell(0,1).text = 'Total Packets'
        table.cell(0,2).text = 'Total Errors'
        table.cell(0,3).text = 'Total Drops'
        table.cell(0,4).text = 'Total Collisions'  
        start_row = 1
        for i, item in enumerate(results['interface_errors']['interfaces']):
            offset = start_row + i
            table.cell(offset, 0).text = item
            table.cell(offset, 1).text = humanize_int(results['interface_errors']['interfaces'][item]['packets_total'])
            table.cell(offset, 2).text = humanize_int(results['interface_errors']['interfaces'][item]['errors_total'])
            table.cell(offset, 3).text = humanize_int(results['interface_errors']['interfaces'][item]['drops_total'])
            table.cell(offset, 4).text = humanize_int(results['interface_errors']['interfaces'][item]['collisions'])

    if problems_found['arp']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['arp']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['arp']['rating'], style='sev'+problems_found['arp']['rating'])
        document.add_paragraph(problems_found['arp']['details'])
        document.add_paragraph()

    if problems_found['irules']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['irules']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['irules']['rating'], style='sev'+problems_found['irules']['rating'])
        document.add_paragraph(problems_found['irules']['details'])
        document.add_paragraph()

        document.add_paragraph('Review the irules that had reported Failures.')
        document.add_paragraph()
        size = len(results['irule_failures'])
        i=0
        print(size)
        table = document.add_table(rows=size+1, cols=4)
        table.style = 'Table Grid'
        table.cell(0,0).text = 'iRule Name'
        table.cell(0,1).text = 'Event'
        table.cell(0,2).text = 'Failures'
        table.cell(0,3).text = 'Total Executions'
        start_row = 1
        for i, item in enumerate(results['irule_failures']):
            offset = start_row + i
            table.cell(offset, 0).text = item['name']
            table.cell(offset, 1).text = item['event']
            table.cell(offset, 2).text = str(item['failures'])
            table.cell(offset, 3).text = str(item['total_executions'])
   
    if problems_found['expired_cert']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['expired_cert']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['expired_cert']['rating'], style='sev'+problems_found['expired_cert']['rating'])
        document.add_paragraph(problems_found['expired_cert']['details'])
        document.add_paragraph()

        document.add_paragraph("Below you can find a list of Certificates Expiring Soon (< 60 days) or already expired: ")

        add_ssl_certs_table_compact_expiring(
            document,
            results["ssl_certs"],
            threshold_days=60
        )
        document.add_paragraph()
        document.add_paragraph()

    if problems_found['disk']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['disk']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['disk']['rating'], style='sev'+problems_found['disk']['rating'])
        document.add_paragraph(problems_found['disk']['details'])
        document.add_paragraph()

        table = document.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.cell(0,0).text = 'Path'
        table.cell(0,1).text = 'Utilization (%)'
        start_row=1
        i=0
        for i, item in enumerate(results['df_h_file']["checks"]):
            offset = start_row + i
            table.cell(offset, 0).text = item['path']
            table.cell(offset, 1).text = str(item["utilization_percent"])

    if problems_found['ntp']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['ntp']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['ntp']['rating'], style='sev'+problems_found['ntp']['rating'])
        document.add_paragraph(problems_found['ntp']['details'])
        document.add_paragraph()

    if problems_found['dns']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['dns']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['dns']['rating'], style='sev'+problems_found['dns']['rating'])
        document.add_paragraph(problems_found['dns']['details'])
        document.add_paragraph()

    if problems_found['cpu']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['cpu']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['cpu']['rating'], style='sev'+problems_found['cpu']['rating'])
        document.add_paragraph(problems_found['cpu']['details'])
        document.add_paragraph()

    if problems_found['http_monitors']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['http_monitors']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['http_monitors']['rating'], style='sev'+problems_found['http_monitors']['rating'])
        document.add_paragraph(problems_found['http_monitors']['details'])
        document.add_paragraph()

        document.add_paragraph('Review the configured HTTP(s) monitors and adjust them as per your application requirements.')
        document.add_paragraph()
        size = len(results['http_monitors'])+len(results['https_monitors'])-1
        table = document.add_table(rows=size+1, cols=6)
        table.style = 'Table Grid'
        table.cell(0,0).text = 'Name'
        table.cell(0,1).text = 'Type'
        table.cell(0,2).text = 'Interval'
        table.cell(0,3).text = 'Timeout'
        table.cell(0,4).text = 'Send String'
        table.cell(0,5).text = 'Receive String'
        start_row = 1
        for i, item in enumerate(results['http_monitors']):
            offset = start_row + i
            table.cell(offset, 0).text = item['name']
            table.cell(offset, 1).text = "HTTP"
            table.cell(offset, 2).text = str(item['interval'])
            table.cell(offset, 3).text = str(item['timeout'])
            table.cell(offset, 4).text = item['send']
            table.cell(offset, 5).text = "N/A" if item.get("recv") is None else str(item.get("recv"))
        start_row=offset
        for i, item in enumerate(results['https_monitors']):
            offset = start_row + i
            table.cell(offset, 0).text = item['name']
            table.cell(offset, 1).text = "HTTPS"
            table.cell(offset, 2).text = str(item['interval'])
            table.cell(offset, 3).text = str(item['timeout'])
            table.cell(offset, 4).text = item['send']
            table.cell(offset, 5).text = "N/A" if item.get("recv") is None else str(item.get("recv"))

    if problems_found['pool_monitor']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['pool_monitor']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['pool_monitor']['rating'], style='sev'+problems_found['pool_monitor']['rating'])
        document.add_paragraph(problems_found['pool_monitor']['details'])
        document.add_paragraph()
        document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
        document.add_paragraph("Below you can find a list of pools (up to 10) that don't have assigned monitors: ")
        count = 0

        for m in results["ltm_pools"]:
            if count>=10:
                continue
            if m.get("monitor") == "none":
                document.add_paragraph("- "+ m.get("pool"))
                count += 1

        document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    

    if problems_found['lb_mode']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['lb_mode']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['lb_mode']['rating'], style='sev'+problems_found['lb_mode']['rating'])
        document.add_paragraph(problems_found['lb_mode']['details'])
        document.add_paragraph()

        document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
        document.add_paragraph("Below you can find a list of pools (up to 10) that have Round Robin as the configured method: ")
        count = 0

        for m in results["ltm_pools"]:
            if count>=10:
                continue
            if m.get("load-balancing-mode") == "round-robin":
                document.add_paragraph("- "+ m.get("pool"))
                count += 1

        document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    

    if problems_found['ssl']['issue']==True:
        document.add_paragraph()
        document.add_heading(problems_found['ssl']['title'], level=3)
        document.add_paragraph('Severity: '+problems_found['ssl']['rating'], style='sev'+problems_found['ssl']['rating'])
        document.add_paragraph(problems_found['ssl']['details'])
        document.add_paragraph()
  
    if len(simplified["diagnostics_simplified"])>0:
        document.add_heading('Security Diagnostics Summary', level=3)
        document.add_paragraph("Please find below all the CVEs that we have identified on your system.")


        # Create table with header row
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "CVE"
        hdr_cells[1].text = "Importance"
        hdr_cells[2].text = "Solution ID"
        hdr_cells[3].text = "Summary"
        hdr_cells[4].text = "FixedInVer"

        for diag in simplified["diagnostics_simplified"]:
            h_importance = diag.get("h_importance", {}).capitalize()
            solution_id = diag.get("solution_id", {})
            #h_cve_ids = diag.get("h_cve_ids", {})
            h_cve_ids = "\n ".join(diag.get("h_cve_ids", []))
            h_header = diag.get("h_header", {})
            fixed_versions = "\n ".join(diag.get("fixedInVersions", []))
            #fixed_versions = diag.get("fixedInVersions", {})
            row_cells = table.add_row().cells
            row_cells[0].text = h_cve_ids
            row_cells[1].text = h_importance
            row_cells[2].text = solution_id
            row_cells[3].text = h_header
            row_cells[4].text = fixed_versions


        set_column_width(table.columns[0], Cm(3.5))
        set_column_width(table.columns[1], Cm(1.5))
        set_column_width(table.columns[2], Cm(2))
        set_column_width(table.columns[3], Cm(9))
        set_column_width(table.columns[4], Cm(0.6))

    for diag in simplified_non_cve["diagnostics_simplified"]:

        if (len(diag["output"])>25):
            diag["output"] = diag["output"][:25]
            diag["output"].append(".............. Output truncated (More info on iHealth) ..............")
            
        if diag["h_importance"] == "CRITICAL":
            header = diag.get("h_header", "")
            # --- SKIP IF IN IGNORE LIST ---

            solution_ids = " and ".join(diag.get("solution_ids", []))
            if  len(diag.get("output"))>0:
                output = "\n\n ".join(diag.get("output", []))
               
            document.add_paragraph()
            document.add_heading(header, level=3)
            document.add_paragraph('Severity: '+diag["h_importance"].capitalize(), style='sev'+diag["h_importance"].capitalize())
            document.add_paragraph(diag["h_summary"])
            document.add_paragraph()
            if  len(diag.get("output"))>0:
                document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
                document.add_paragraph(output)
                document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    
            document.add_paragraph()

            if solution_ids != "":
                document.add_paragraph("Please review the solution article for more information: " + solution_ids)
                document.add_paragraph()

    for diag in simplified_non_cve["diagnostics_simplified"]:

        if diag["h_importance"] == "HIGH":
            header = diag.get("h_header", "")

            solution_ids = " and ".join(diag.get("solution_ids", []))
            if  len(diag.get("output"))>0:
                output = "\n\n ".join(diag.get("output", []))

            document.add_paragraph()
            document.add_heading(header, level=3)
            document.add_paragraph('Severity: '+diag["h_importance"].capitalize(), style='sev'+diag["h_importance"].capitalize())
            document.add_paragraph(diag["h_summary"])
            document.add_paragraph()
            if  len(diag.get("output"))>0:
                document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
                document.add_paragraph(output)
                document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    
            document.add_paragraph()

            if solution_ids != "":
                document.add_paragraph("Please review the solution article for more information: " + solution_ids)
                document.add_paragraph()

    for diag in simplified_non_cve["diagnostics_simplified"]:

        if diag["h_importance"] == "MEDIUM":
            header = diag.get("h_header", "")

            solution_ids = " and ".join(diag.get("solution_ids", []))
            if  len(diag.get("output"))>0:
                output = "\n\n ".join(diag.get("output", []))
                
            document.add_paragraph()
            document.add_heading(header, level=3)
            document.add_paragraph('Severity: '+diag["h_importance"].capitalize(), style='sev'+diag["h_importance"].capitalize())
            document.add_paragraph(diag["h_summary"])
            document.add_paragraph()
            if  len(diag.get("output"))>0:
                document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
                document.add_paragraph(output)
                document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    
            document.add_paragraph()

            if solution_ids != "":
                document.add_paragraph("Please review the solution article for more information: " + solution_ids)
                document.add_paragraph()

    for diag in simplified_non_cve["diagnostics_simplified"]:

        if diag["h_importance"] == "LOW":
            header = diag.get("h_header", "")

            solution_ids = " and ".join(diag.get("solution_ids", []))
            if  len(diag.get("output"))>0:
                output = "\n\n ".join(diag.get("output", []))
                
            document.add_paragraph()
            document.add_heading(header, level=3)
            document.add_paragraph('Severity: '+diag["h_importance"].capitalize(), style='sev'+diag["h_importance"].capitalize())
            document.add_paragraph(diag["h_summary"])
            document.add_paragraph()
            if  len(diag.get("output"))>0:
                document.add_paragraph("--------------------------------------------------------------  Additional details  --------------------------------------------------------------")
                document.add_paragraph(output)
                document.add_paragraph("------------------------------------------------------------------------------------------------------------------------------------------------------")    
            document.add_paragraph()

            if solution_ids != "":
                document.add_paragraph("Please review the solution article for more information: " + solution_ids)
                document.add_paragraph()

    today = datetime.date.today()
    today_str = today.isoformat()

    output_dir = Path("reports")
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / f"F5 LTM - Config Review - {customer_name} - {results['list_cm_device']['hostname']} - {today_str}.docx"

    document.save(output_file)
    print(f"Saved Word: {output_file}")

def humanize_int(value: int):
    """
    Convert an integer to a human-readable format like:
      1200 -> '1.2K'
      2000000 -> '2.0M'
      3500000000 -> '3.5G'
    Supports K, M, G, T, P.
    """
    if value is None:
        return "0"

    num = float(value)

    for unit in ["", "K", "M", "G", "T", "P"]:
        if num < 1000:
            return f"{num:.1f}{unit}" if unit else str(int(num))
        num /= 1000.0

    # fallback for insanely large numbers
    return f"{num:.1f}E"   # Exa (beyond Peta)

def merge_json_objects(**objects) -> dict:
    """
    Merge multiple JSON-like Python objects into a single dict.
    Keys are provided by the caller.
    """
    merged_data = {}

    for key, value in objects.items():
        if value is None:
            continue
        merged_data[key] = value

    return merged_data

def test_access(token):
    try:
        r = requests.get(COMMAND_BASE, headers={"Authorization": f"Bearer {token}", "Accept": "application/json"}, timeout=15)
        return r.status_code == 200
    except Exception as e:
        print({e})
        return False

def get_token(CLIENT_ID, CLIENT_SECRET, TOKEN_URL):
    
    print("🔐 Requesting new token...")
    resp = requests.post(
        TOKEN_URL,
        headers={
            "Accept": "application/json",
            "Content-Type": "application/x-www-form-urlencoded",
        },
        data={"grant_type": "client_credentials", "scope": "ihealth"},
        auth=(CLIENT_ID, CLIENT_SECRET),
        timeout=30,
    )
    resp.raise_for_status()
    token_data = resp.json()
    print("✅ New token acquired.")

    if not test_access(token_data.get("access_token")):
        raise SystemExit("❌ Token obtained but access failed.")
    return token_data.get("access_token")

def decode_output(b64text: str) -> str:
    cleaned = re.sub(r"[%\n]", "", b64text or "")
    # fix padding just in case
    pad = (-len(cleaned)) % 4
    if pad:
        cleaned += "=" * pad
    # decode as text (ignore undecodable bytes)
    return base64.b64decode(cleaned).decode("utf-8", errors="ignore")

def fetch_with_retry(url: str, headers: dict, timeout: int = 60, retries: int = 3, delay: int = 5):
    """Fetch a URL with retry logic. Retry on ANY exception or non-200 status."""
    for attempt in range(1, retries + 1):
        try:
            r = requests.get(url, headers=headers, timeout=timeout)
            if r.status_code == 200 or r.status_code == 404:
                return r     
            print(f"⚠️ Attempt {attempt}/{retries} failed (HTTP {r.status_code}). Retrying in {delay}s...")
        except Exception as e:
            print(f"⚠️ Attempt {attempt}/{retries} failed ({e}). Retrying in {delay}s...")
        time.sleep(delay)
    print(f"❌ Failed after {retries} attempts. Skipping.")
    return None

def parse_partition_names(text: str) -> list[str]:
    """
    Extract partition names from 'auth partition <name> {' blocks.
    """
    return re.findall(r'\bauth\s+partition\s+([^\s]+)\s*\{', text)

def get_partitions(token, headers):
    command ="list /auth partition all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 

    r = fetch_with_retry(url, headers)
    if not r:
        print("⚠️ Partitions failed to be retrieved.")
        exit(1)  # skip this command
    data = r.json()

    # Try .[0].output then .output
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")

    if not output:
        print(f"⚠️ No partitions found output for:")
        exit(1)

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for Partitions: {e}")
        exit(1)

    partitions = parse_partition_names(decoded) if decoded else []
    print(f"Found {len(partitions)} Partitions.")
    return partitions

def get_security_profiles(token, headers):
    """
    Extract security profile names from the bigip.conf on all partitions.
    profile_type: e.g. "security dos profile", "security bot-defense profile", "asm policy"
    """
    partitions=get_partitions(token, headers)
    dos=[]
    asm=[]
    bot=[]

    for partition in partitions:
        print(f"Processing partition: {partition}")
        
        if partition == "Common":
            partition_name = "config/bigip.conf"
            base64_partition_name = base64.b64encode(partition_name.encode("utf-8")).decode("utf-8").rstrip("=")
        else:
            partition_name = "config/partitions/"+partition+"/bigip.conf"
            base64_partition_name = base64.b64encode(partition_name.encode("utf-8")).decode("utf-8").rstrip("=")


        detail_url = FILES_URL + "/" + base64_partition_name  # tmsh list sys partition"
        print(detail_url)
        
        r = fetch_with_retry(detail_url, headers=headers, timeout=60, retries=3, delay=5)
        if r.status_code == 404:
            print(f"⚠️ Bigip.conf File Not Found for Partition {partition}.")
            continue  # skip this command
        if not r:
            print(f"⚠️ Command failed for partition {partition}.")
            continue  # skip this command

        dos.extend(parse_profiles(r.text,"security dos profile"))
        bot.extend(parse_profiles(r.text,"security bot-defense profile"))
        asm.extend(parse_profiles(r.text,"asm policy"))


    print(f"Found {len(dos)} DOS Profiles")
    print(f"Found {len(bot)} Bot Defense Profiles")
    print(f"Found {len(asm)} ASM Policies")
    return dos, bot, asm

def get_web_acceleration(token, headers):
    """
    Collect and extract webacceleration profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile web-acceleration all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing Web Acceleration Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None


    result= parse_profiles(decoded,"ltm profile web-acceleration")

    print(f"Found {len(result)} Web Acceleration Profiles")

    return result

def get_udp_profiles(token, headers):
    """
    Collect and extract UDP profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile udp all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing UDP Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm profile udp")

    print(f"Found {len(result)} UDP Profiles")

    return result

def get_tcp_profiles(token, headers):
    """
    Collect and extract TCP profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile tcp all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing TCP Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm profile tcp")

    print(f"Found {len(result)} TCP Profiles")

    return result

def get_server_ssl_profiles(token, headers):
    """
    Collect and extract Server SSL profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile server-ssl all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing Server SSL Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm profile server-ssl")

    print(f"Found {len(result)} Server SSL Profiles")

    return result

def get_http2_profiles(token, headers):
    """
    Collect and extract HTTP2 profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile http2 all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing HTTP2 Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm profile http")

    print(f"Found {len(result)} HTTP2 Profiles")

    return result

def get_http_compression_profiles(token, headers):
    """
    Collect and extract HTTP compression profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile http-compression all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing HTTP Compression Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm profile http-compression")

    print(f"Found {len(result)} HTTP Compression Profiles")

    return result

def get_http_profiles(token, headers):
    """
    Collect and extract HTTP profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile http all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing HTTP Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_http_profiles(decoded)

    print(f"Found {len(result)} HTTP Profiles")

    return result

def get_client_ssl_profiles(token, headers):
    """
    Collect and extract Client SSL profile names from the ihealth.
    """
    result=[]
    command ="list /ltm profile client-ssl all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing Client SSL Profiles")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_client_ssl_profiles(decoded)

    print(f"Found {len(result)} Client SSL Profiles")

    return result

def get_sys_hardware(token, headers):
    """
    Collect and extract sys_hardware from the ihealth.
    """
    
    command ="show /sys hardware"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing sys hardware")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_sys_hardware(decoded)

    print(f"Found {len(result)} sys hardware items")

    return result

def get_sync_status(token, headers):
    """
    Collect and extract sync-status from the ihealth.
    """
    
    command ="show /cm sync-status"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing sync-status")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_sync_status(decoded)

    print(f"Found {len(result)} sync status items")

    return result

def get_ssl_certificates(token, headers):
    """
    Analyze "show public ssl certificates" output to extract certificate details.
    """
    
    command ="Public SSL Certificates"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing SSL certificates")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_ssl_certificates(decoded)

    print(f"Found {len(result)} SSL certificates")

    return result

def get_sys_provision(token, headers):
    """
    Analyze "list sys provision all-properties" output to extract provisioned modules.
    """
    
    command ="list /sys provision all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing sys provision")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_sys_provision(decoded)

    print(f"Found {len(result)} provisioned modules")

    return result

def get_ltm_pools(token, headers):
    """
    Analyze "list ltm pool all-properties" output to extract pool details.
    """
    
    command ="list /ltm pool all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing LTM pools")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_ltm_pools(decoded)

    print(f"Found {len(result)} LTM pools")

    return result

def get_ltm_virtual(token, headers):
    """
    Analyze "list ltm virtual all-properties" output to extract virtual details.
    """
    
    command ="list /ltm virtual all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing LTM virtuals")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_ltm_virtuals(decoded)

    print(f"Found {len(result)} LTM virtuals")

    return result

def get_self_ips(token, headers):
    """
    Analyze "show running-config /net self" output to extract self IP details.
    """
    
    command ="show running-config /net self"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing Self IPs")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_self_ips(decoded)

    print(f"Found {len(result)} LTM self IPs")

    return result

def get_show_cm_device(token, headers):
    """
    Analyze "show /cm device" output to extract device details.
    """
    
    command ="show /cm device"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /cm device")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_show_cm_device(decoded)

    print(f"Found {len(result)} CM devices")

    return result

def get_list_cm_device(token, headers):
    """
    Analyze "list /cm device" output to extract device details.
    """
    
    command ="list /cm device"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing List CM Devices")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_list_cm_device(decoded)

    print(f"Found {len(result)} CM devices")

    return result

def get_icmp_monitors(token, headers):
    """
    Analyze "list /ltm monitor icmp all-properties" output to extract ICMP monitor details.
    """
    
    command = "list /ltm monitor icmp all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing ICMP Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_other_monitors(decoded)

    print(f"Found {len(result)} ICMP Monitors")

    return result

def get_udp_monitors(token, headers):
    """
    Analyze "list /ltm monitor udp all-properties" output to extract UDP monitor details.
    """
    
    command = "list /ltm monitor udp all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing UDP Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_other_monitors(decoded)

    print(f"Found {len(result)} UDP Monitors")

    return result

def get_tcp_monitors(token, headers):
    """
    Analyze "list /ltm monitor tcp all-properties" output to extract TCP monitor details.
    """
    
    command = "list /ltm monitor tcp all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing TCP Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_other_monitors(decoded)

    print(f"Found {len(result)} TCP Monitors")

    return result

def get_tcp_half_open_monitors(token, headers):
    """
    Analyze "list /ltm monitor tcp-half-open all-properties" output to extract TCP Half Open monitor details.
    """
    
    command = "list /ltm monitor tcp-half-open all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing TCP Half Open Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_other_monitors(decoded)

    print(f"Found {len(result)} TCP Half Open Monitors")

    return result

def get_http_monitors(token, headers):
    """
    Analyze "list /ltm monitor http all-properties" output to extract HTTP monitor details.
    """
    
    command = "list /ltm monitor http all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing HTTP Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_http_monitors(decoded)

    print(f"Found {len(result)} HTTP Monitors")

    return result

def get_https_monitors(token, headers):
    """
    Analyze "list /ltm monitor https all-properties" output to extract HTTPS monitor details.
    """
    
    command = "list /ltm monitor https all-properties"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing HTTPS Monitors")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_http_monitors(decoded)

    print(f"Found {len(result)} HTTPS Monitors")

    return result

def get_df_h_file(token, headers):
    """
    Analyze "df -h" output to extract disk usage details.
    """
    
    command = "df -h"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing df -h output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_df_h_file(decoded)

    print(f"Found {len(result)} df -h entries")

    return result

def get_license_file(token, headers):
    """
    Analyze "show /sys license detail" output to extract license details.
    """
    
    command = "show /sys license detail"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /sys license detail output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_license_file(decoded)

    print(f"Finished processing license entries")

    return result

def get_interface_errors(token, headers):
    """
    Analyze "show /net interface all-properties -hidden" output to extract interface error details.
    """
    
    command = "show /net interface all-properties -hidden"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /net interface all-properties -hidden output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    per_iface= parse_interface_table(decoded)
    summary = build_interface_summary(per_iface)
    result = {
        "interfaces": per_iface,
        "summary": summary,
    }
    print(f"Finished processing interface errors")

    return result

def get_http_profile_global(token, headers):
    """
    Analyze "show /ltm profile http global" output to extract HTTP stats.
    """
    
    command = "show /ltm profile http global"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /ltm profile http global output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_http_profile_global(decoded)

    print(f"Finished processing http profile global entries")

    return result

def get_irule_failures(token, headers):
    """
    Analyze "show /ltm rule all" output to extract IRule stats.
    """
    
    command = "show /ltm rule all"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /ltm rule all output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    irules_stats, failures  = parse_failures_execs_by_event(decoded)

    print(f"Finished processing irule stats")

    return irules_stats, failures

def get_irules(token, headers):
    """
    Analyze "list /ltm rule" output to extract IRule names.
    """
    
    command = "list /ltm rule"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing list /ltm rule output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_profiles(decoded,"ltm rule")

    print(f"Finished processing irules names")

    return result

def get_client_ssl_stats(token, headers):
    """
    Analyze "show /ltm profile client-ssl global" output to extract SSL stats.
    """
    
    command = "show /ltm profile client-ssl global"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /ltm profile client-ssl global output")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_client_ssl_stats(decoded)

    print(f"Finished processing client-ssl profile global entries")

    return result

def get_tmctl_stats(token, headers):
    """
    Collect and extract tmctl (global stats).
    """
    result=[]
    command ="tmctl (global stats)"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing tmctl (global stats)")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_tmctl_stats(decoded)

    print(f"Finished processing tmctl (global stats)")

    return result

def get_host_info(token, headers):
    """
    Collect and extract show /sys host-info global.
    """
    result=[]
    command ="show /sys host-info global"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /sys host-info global")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_host_info(decoded)

    print(f"Finished processing show /sys host-info global")

    return result

def get_sysdb_file(token, headers):
    """
    Collect and extract list /sys db all-properties.
    """
    result=[]
    command ="list /sys db all-properties (non-default values)"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing list /sys db all-properties")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_sysdb_file(decoded)

    print(f"Finished processing list /sys db all-properties")

    return result

def get_version_file(token, headers):
    """
    Collect and extract show /sys version detail.
    """
    result=[]
    command ="show /sys version detail"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /sys version detail")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_version_file(decoded)

    print(f"Finished processing show /sys version detail")

    return result

def get_arp(token, headers):
    """
    Collect and extract show /net arp
    """
    result=[]
    command ="show /net arp"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /net arp")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_arp(decoded)

    print(f"Finished processing show /net arp")

    return result

def get_virtual_server_traffic(token, headers):
    """
    Collect and extract  Virtual Server Traffic
    """
    result=[]
    command ="Virtual Server Traffic"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing show /net arp")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    result= parse_vs_stats(decoded)

    print(f"Finished processing  Virtual Server Traffic")

    return result

def get_date(token, headers):
    """
    Collect the date.
    """
    result=[]
    command ="date"
    url = COMMAND_BASE + hashlib.sha1(command.encode("utf-8")).hexdigest() 
    print("Processing date")
    r = fetch_with_retry(url, headers=headers, timeout=60, retries=3, delay=5)
    if r.status_code == 404:
        print(f"⚠️ Command ({command}) Not Found.")
        return None  
    if not r:
        print(f"⚠️ Command  ({command}) failed.")
        return None 
    
    data = r.json()
    output = None
    if isinstance(data, list) and data and isinstance(data[0], dict):
        output = data[0].get("output")
    if not output:
        print(f"⚠️ No output for: {command}")
        return None 

    try:
        decoded = decode_output(output)
    except Exception as e:
        print(f"❌ Base64 decode failed for: {command} ({e})")
        return None

    date_str= parse_date(decoded)
    result = {"system_date": date_str}

    print(f"Finished processing date")

    return result

def _profiles_from_dict_list(items, key="profile"):
    return {
        item.get(key)
        for item in (items or [])
        if isinstance(item, dict) and item.get(key)
    }

def create_vs_summary(virtuals,dos, bot, tcp_profiles, udp_profiles, http_profiles, http2_profiles, compression_profiles, client_ssl_profiles, server_ssl_profiles, web_acceleration_profiles):

    http_set = _profiles_from_dict_list(http_profiles)
    clientssl_set = _profiles_from_dict_list(client_ssl_profiles)

    vs_details = []
    count_vs=0
    count_vs_dos=0
    count_vs_bot=0
    count_vs_web_acceleration=0
    count_vs_udp=0
    count_vs_tcp=0
    count_vs_compression=0
    count_vs_server_ssl=0
    count_vs_http=0
    count_vs_client_ssl=0
    count_vs_asm=0
    for v in virtuals:
        v_profiles = set(v.get("profiles") or [])
        v_policies = v.get("policies") or []

        # ASM check: substring match in policies
        has_asm = any(
            "asm_auto_l7_policy" in policy
            for policy in v_policies
            if isinstance(policy, str)
        )
        item = {
            "virtual": v.get("virtual"),
            "partition": v.get("partition"),

            "has_dos": bool(v_profiles & set(dos)),
            "has_bot": bool(v_profiles & set(bot)),
            "has_web_acceleration": bool(v_profiles & set(web_acceleration_profiles)),
            "has_udp": bool(v_profiles & set(udp_profiles)),
            "has_tcp": bool(v_profiles & set(tcp_profiles)),
            "has_compression": bool(v_profiles & set(compression_profiles)),
            "has_server_ssl": bool(v_profiles & set(server_ssl_profiles)),
            "has_http": bool(v_profiles & http_set),
            "has_client_ssl": bool(v_profiles & clientssl_set),
            "has_asm": has_asm
        }
        count_vs += 1
        if item["has_dos"]:
            count_vs_dos += 1
        if item["has_bot"]:
            count_vs_bot += 1
        if item["has_web_acceleration"]:
            count_vs_web_acceleration += 1
        if item["has_udp"]:
            count_vs_udp += 1
        if item["has_tcp"]:
            count_vs_tcp += 1
        if item["has_compression"]:
            count_vs_compression += 1
        if item["has_server_ssl"]:
            count_vs_server_ssl += 1
        if item["has_http"]:
            count_vs_http += 1
        if item["has_client_ssl"]:
            count_vs_client_ssl += 1
        if item["has_asm"]:
            count_vs_asm += 1

        vs_details.append(item)


    summary = {
        "count_vs": count_vs,
        "count_vs_dos": count_vs_dos,
        "count_vs_bot": count_vs_bot,
        "count_vs_web_acceleration": count_vs_web_acceleration,
        "count_vs_udp": count_vs_udp,
        "count_vs_tcp": count_vs_tcp,
        "count_vs_compression": count_vs_compression,
        "count_vs_server_ssl": count_vs_server_ssl,
        "count_vs_http": count_vs_http,
        "count_vs_client_ssl": count_vs_client_ssl,
        "count_vs_asm": count_vs_asm
    }

    return vs_details, summary


parser = argparse.ArgumentParser(description="Generate LTM reports from a qkview")
parser.add_argument(
    "--customer",
    "-c",
    required=True,
    help="Customer name to be used in the report",
)
parser.add_argument(
    "--qkview_id",
    "-q",
    required=True,
    help="QKVIEW ID to be used in the report",
)
parser.add_argument(
    "--client_id",
    "-cid",
    required=False,
    help="CLIENT ID to be used for iHealth",
)
parser.add_argument(
    "--client_secret",
    "-s",
    required=False,
    help="CLIENT SECRET to be used for iHealth",
)

args = parser.parse_args()
QKVIEW_ID = args.qkview_id

if args.client_id == None:
    try:
        CLIENT_ID = os.environ["CLIENT_ID"]
    except KeyError:
        raise RuntimeError("CLIENT_ID is neither set as an environment variable or send through the command line.")
else:
    CLIENT_ID = args.client_id

print(args.client_id)

if args.client_secret == None:
    try:
        CLIENT_SECRET = os.environ["CLIENT_SECRET"]
    except KeyError:
        raise RuntimeError("CLIENT_SECRET is neither set as an environment variable or send through the command line.")
else:
    CLIENT_ID = args.client_secret

customer_name = args.customer

# --- CONFIGURATION ---

TOKEN_URL = "https://identity.account.f5.com/oauth2/ausp95ykc80HOU7SQ357/v1/token"



LIST_URL = f"https://ihealth-api.f5.com/qkview-analyzer/api/qkviews/{QKVIEW_ID}/commands"
COMMAND_BASE = f"https://ihealth-api.f5.com/qkview-analyzer/api/qkviews/{QKVIEW_ID}/commands/"
DIAG_URL = f"https://ihealth2-api.f5.com/qkview-analyzer/api/qkviews/{QKVIEW_ID}/diagnostics"
FILES_URL = f"https://ihealth2-api.f5.com/qkview-analyzer/api/qkviews/{QKVIEW_ID}/files"
# ----------------------




print("\n🏁 Start.")

token = get_token(CLIENT_ID, CLIENT_SECRET, TOKEN_URL)
headers = {"Authorization": f"Bearer {token}", "Accept": "application/vnd.f5.ihealth.api+json"}

resp = fetch_with_retry(DIAG_URL, headers)
if not resp or resp.status_code != 200:
    diagnostics = {}
else:
    diagnostics = resp.json()

dos, bot, asm = get_security_profiles(token, headers)
irule_stats, irule_failures = get_irule_failures(token, headers)
web_acceleration=get_web_acceleration(token, headers)
udp_profiles=get_udp_profiles(token, headers)
tcp_profiles=get_tcp_profiles(token, headers)
server_ssl=get_server_ssl_profiles(token, headers)
http2=get_http2_profiles(token, headers)
compression=get_http_compression_profiles(token, headers)
http=get_http_profiles(token, headers)
client_ssl=get_client_ssl_profiles(token, headers)
ltm_pools=get_ltm_pools(token, headers)
ltm_virtuals=get_ltm_virtual(token, headers)
virtual_details, vs_sum = create_vs_summary(ltm_virtuals, dos, bot, tcp_profiles, udp_profiles, http, http2, compression, client_ssl, server_ssl, web_acceleration)

results = merge_json_objects(
    dos=dos,
    bot=bot,
    asm=asm,
    web_acceleration=web_acceleration,
    udp_profiles=udp_profiles,
    tcp_profiles=tcp_profiles,
    server_ssl=server_ssl,
    http2=http2,
    compression=compression,
    http=http,
    client_ssl=client_ssl,
    sys_hardware=get_sys_hardware(token, headers),
    sync_status=get_sync_status(token, headers),
    ssl_certs=get_ssl_certificates(token, headers),
    sys_provision=get_sys_provision(token, headers),
    ltm_pools=ltm_pools,
    ltm_virtuals=ltm_virtuals,
    self_ips=get_self_ips(token, headers),
    show_cm_device=get_show_cm_device(token, headers),
    list_cm_device=get_list_cm_device(token, headers),
    icmp_monitors=get_icmp_monitors(token, headers),
    udp_monitors=get_udp_monitors(token, headers),
    tcp_monitors=get_tcp_monitors(token, headers),
    tcp_half_open_monitors=get_tcp_half_open_monitors(token, headers),
    http_monitors=get_http_monitors(token, headers),
    https_monitors=get_https_monitors(token, headers),
    df_h_file=get_df_h_file(token, headers),
    license_file=get_license_file(token, headers),
    interface_errors=get_interface_errors(token, headers),
    http_profile_global=get_http_profile_global(token, headers),
    irule_stats=irule_stats,
    irule_failures=irule_failures,
    irule_names=get_irules(token, headers),
    client_ssl_stats=get_client_ssl_stats(token, headers),
    tmctl_stats=get_tmctl_stats(token, headers),
    host_info=get_host_info(token, headers),
    sysdb_file=get_sysdb_file(token, headers),
    version_file=get_version_file(token, headers),
    arp_table=get_arp(token, headers),
    date_info=get_date(token, headers),
    virtual_details=virtual_details,
    virtual_summary=vs_sum,
    vs_traffic=get_virtual_server_traffic(token, headers)
    )

build_excel(results, customer_name)
build_word(results, diagnostics, customer_name)


print("\n🏁 All done.")


